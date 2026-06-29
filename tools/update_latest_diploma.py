#!/usr/bin/env python3

import argparse
import copy
import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader


SOURCE = Path(
    "/Users/limon4ik/Downloads/"
    "Диплом_Жунусов_обновленное_введение_аннотация.docx"
)

FONT_REGULAR = "/System/Library/Fonts/Supplemental/Times New Roman.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"


def image_font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def find_heading(doc, exact_text):
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if "Heading" in style_name and paragraph.text.strip() == exact_text:
            return paragraph
    raise ValueError(f"Heading not found: {exact_text}")


def replace_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    return paragraph


def set_paragraph(doc, prefix, text):
    return replace_paragraph_text(find_paragraph(doc, prefix), text)


def insert_after(paragraph, text):
    new_p = copy.deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)

    source_run = paragraph.runs[0]._r if paragraph.runs else None
    if source_run is not None:
        new_run = copy.deepcopy(source_run)
        for child in list(new_run):
            if child.tag != qn("w:rPr"):
                new_run.remove(child)
    else:
        from docx.oxml import OxmlElement

        new_run = OxmlElement("w:r")

    from docx.oxml import OxmlElement

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    new_p.append(new_run)
    paragraph._p.addnext(new_p)

    from docx.text.paragraph import Paragraph

    return Paragraph(new_p, paragraph._parent)


def find_table(doc, first_cell, columns):
    for table in doc.tables:
        if len(table.columns) == columns and table.cell(0, 0).text.strip() == first_cell:
            return table
    raise ValueError(f"Table not found: {first_cell}/{columns}")


def set_run_font(run, size=12, bold=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:ascii"), "Times New Roman"
    )
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:hAnsi"), "Times New Roman"
    )
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), "Times New Roman"
    )
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, size=12, align=None):
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            set_run_font(run, size=size)


def append_row(table, values, size=12):
    new_tr = copy.deepcopy(table.rows[-1]._tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
    return row


def replace_row(table, row_index, values, size=12):
    row = table.rows[row_index]
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value, size=size)


def draw_grouped_chart(
    path,
    size,
    title,
    categories,
    series,
    y_max,
    y_label,
    decimals=0,
):
    width, height = size
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)

    left, right = 170, width - 100
    top, bottom = 165, height - 175
    plot_width = right - left
    plot_height = bottom - top

    draw.text((80, 45), title, fill="black", font=image_font(50, bold=True))

    tick_font = image_font(28)
    label_font = image_font(27)
    value_font = image_font(24, bold=True)
    legend_font = image_font(27)

    for tick in range(6):
        value = y_max * tick / 5
        y = bottom - plot_height * tick / 5
        draw.line((left, y, right, y), fill="#d0d0d0", width=2)
        label = f"{value:.0f}" if decimals == 0 else f"{value:.1f}"
        bbox = draw.textbbox((0, 0), label, font=tick_font)
        draw.text(
            (left - 22 - (bbox[2] - bbox[0]), y - 15),
            label,
            fill="#333333",
            font=tick_font,
        )

    draw.line((left, top, left, bottom), fill="#222222", width=4)
    draw.line((left, bottom, right, bottom), fill="#222222", width=4)
    draw.text((35, (top + bottom) / 2 - 20), y_label, fill="#333333", font=label_font)

    colors = ["#3f6da8", "#c74f55"]
    group_width = plot_width / len(categories)
    bar_width = min(126, group_width * 0.26)
    gap = 18

    for index, category in enumerate(categories):
        center = left + group_width * (index + 0.5)
        for series_index, (_, values) in enumerate(series):
            value = values[index]
            x0 = center + (series_index - 0.5) * (bar_width + gap)
            x1 = x0 + bar_width
            y0 = bottom - (value / y_max) * plot_height
            draw.rectangle((x0, y0, x1, bottom), fill=colors[series_index], outline="#222")
            value_text = (
                f"{value:.{decimals}f}" if decimals else f"{int(round(value))}"
            )
            bbox = draw.textbbox((0, 0), value_text, font=value_font)
            draw.text(
                ((x0 + x1 - (bbox[2] - bbox[0])) / 2, y0 - 34),
                value_text,
                fill="black",
                font=value_font,
            )

        bbox = draw.textbbox((0, 0), category, font=label_font)
        draw.text(
            (center - (bbox[2] - bbox[0]) / 2, bottom + 28),
            category,
            fill="black",
            font=label_font,
        )

    legend_y = height - 72
    legend_x = left
    for index, (name, _) in enumerate(series):
        x = legend_x + index * 470
        draw.rectangle((x, legend_y, x + 38, legend_y + 28), fill=colors[index], outline="#222")
        draw.text((x + 55, legend_y - 3), name, fill="black", font=legend_font)

    image.save(path)


def patch_media(docx_path, replacements):
    with tempfile.TemporaryDirectory(prefix="diploma_patch_") as temp_name:
        temp_path = Path(temp_name) / docx_path.name
        with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
            temp_path, "w", zipfile.ZIP_DEFLATED
        ) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename in replacements:
                    data = Path(replacements[item.filename]).read_bytes()
                target.writestr(item, data)
        temp_path.replace(docx_path)


def build_document(output):
    doc = Document(SOURCE)

    set_paragraph(
        doc,
        "В работе спроектирована",
        "В работе спроектирована и реализована система централизованного "
        "логирования распределённых приложений для тренажёрной системы. Сервер "
        "принимает сигналы и события через WebSocket, сохраняет их пакетами в "
        "ClickHouse, предоставляет REST API для чтения, архивирует завершённые "
        "сессии в JSON Lines с gzip-сжатием и удаляет данные по запросу. "
        "Реализация выполнена на Go и развёртывается в Docker Compose. Часовое "
        "испытание подтвердило скорость 4999 сигналов в секунду без потерь, а "
        "архивирование большой сессии заняло 19,94–29,12 секунды.",
    )
    set_paragraph(
        doc,
        "The thesis designs",
        "The thesis designs and implements a centralized logging system for "
        "distributed applications intended for a training simulator. The server "
        "receives signals and events through WebSocket, writes them to ClickHouse "
        "in batches, provides REST API access, creates gzip-compressed JSON Lines "
        "archives, and deletes data on request. A one-hour test confirmed 4,999 "
        "signals per second without loss, while archiving a large session took "
        "19.94–29.12 seconds.",
    )

    set_paragraph(
        doc,
        "HTTP-сервер задаёт",
        "HTTP-сервер задаёт ReadTimeout и WriteTimeout по 30 секунд. Middleware "
        "с ограничением 60 секунд применяется к health-check и REST-маршрутам, "
        "но исключён из WebSocket-маршрута. После upgrade долгоживущее соединение "
        "управляется библиотекой gorilla/websocket и механизмом ping/pong. "
        "Разделение было выполнено после высокого профиля, в котором общий HTTP "
        "timeout отменял финальный flush на 60-й секунде.",
    )
    set_paragraph(
        doc,
        "Ping отправляется",
        "Ping отправляется из отдельной горутины. Нагрузочный клиент постоянно "
        "читает соединение, поэтому библиотека обрабатывает служебные кадры и "
        "отправляет pong в длительных сессиях. Если сервер начнёт передавать "
        "прикладные сообщения клиенту, конкурентную запись в websocket.Conn "
        "потребуется синхронизировать.",
    )
    set_paragraph(
        doc,
        "Утилита cmd/loadtest",
        "Утилита cmd/loadtest создаёт отдельную горутину и WebSocket-соединение "
        "для каждой сессии. Сигналы отправляются пакетами по 50 сообщений. "
        "Частота задаётся флагами signals-per-sec, events-per-sec, sessions и "
        "duration. В Makefile добавлен профиль loadtest-hour: 10 сессий, по 500 "
        "сигналов и 20 событий в секунду на сессию в течение одного часа.",
    )
    set_paragraph(
        doc,
        "Генератор фиксирует",
        "Генератор фиксирует сетевые сбои отправки и отдельно выводит клиентские "
        "счётчики. Фоновый цикл чтения WebSocket необходим для обработки ping/pong "
        "в длительных соединениях. Для полноценной проверки счётчики дополнительно "
        "сравниваются с числом строк в ClickHouse.",
    )
    set_paragraph(
        doc,
        "Реализация соответствует",
        "Реализация соответствует спроектированным слоям и успешно собирается "
        "командой go test ./.... Раздельная политика timeout и обработка ping/pong "
        "проверены часовым нагрузочным профилем. Оставшиеся направления развития "
        "связаны с возвратом буфера при произвольной ошибке репозитория, "
        "объединением списков активных буферов и постраничным чтением архива.",
    )

    archive_anchor = set_paragraph(
        doc,
        "Архивирование сигналов и событий выполняется",
        "Архивирование сигналов и событий выполняется последовательно. HTTP-клиент "
        "ожидает сброса буферов, чтения записей из ClickHouse и полного создания "
        "обоих gzip-файлов. Только после этого обработчик возвращает HTTP 200 со "
        "статусом archived.",
    )
    archive_anchor = insert_after(
        archive_anchor,
        "После создания файлов сервис немедленно меняет статус в памяти, а затем "
        "запускает фоновую goroutine для обновления статуса в ClickHouse и удаления "
        "партиций signal_logs и event_logs. Клиент эту стадию не ожидает и "
        "отдельного уведомления о её завершении не получает; ошибки записываются "
        "только в серверный журнал.",
    )
    insert_after(
        archive_anchor,
        "REST-маршрут ограничен timeout 60 секунд. Поэтому при дальнейшем росте "
        "объёма синхронное формирование архива может быть отменено до ответа. Для "
        "крупных сессий целесообразно запускать архивирование как фоновое задание и "
        "возвращать клиенту идентификатор операции с состояниями queued, running, "
        "completed и failed.",
    )

    set_paragraph(
        doc,
        "Использовались три профиля",
        "Использовались четыре профиля из Makefile. Лёгкий профиль проверяет "
        "базовый сценарий. Средний увеличивает суммарный поток до 1500 сигналов в "
        "секунду. Высокий запускает 10 соединений и целевой поток 5000 сигналов в "
        "секунду в течение 60 секунд. Часовой профиль повторяет высокую нагрузку "
        "3600 секунд и используется как регрессионная проверка исправления.",
    )
    set_paragraph(
        doc,
        "После каждого теста выполнялся",
        "После каждого теста выполнялся запрос количества строк по каждой сессии. "
        "Высокий профиль локализовал потерю последней порции сессии 1000 из-за "
        "общего HTTP timeout. После изменения маршрутизации и обработки ping/pong "
        "часовой профиль проверил исправление на длительности, в 60 раз большей.",
    )
    set_paragraph(
        doc,
        "Результаты трёх профилей",
        "Результаты четырёх профилей приведены в таблице 5.2.",
    )

    load_table = find_table(doc, "Профиль", 6)
    append_row(
        load_table,
        ["1 час", "10", "≈18 млн", "4 999", "≈18 млн", "0"],
        size=9,
    )

    result_anchor = set_paragraph(
        doc,
        "До завершения профиль выполнялся",
        "До завершения исходный высокий профиль выполнялся стабильно. Расхождение "
        "сформировалось только при финальном сбросе одной сессии и относится не к "
        "пропускной способности ClickHouse, а к политике завершения WebSocket.",
    )
    result_anchor = insert_after(
        result_anchor,
        "После точечного исправления выполнен четвёртый профиль. За 3600 секунд "
        "отправлено 17 995 850 сигналов и 719 665 событий. Средняя скорость "
        "составила 4999 сигналов и около 200 событий в секунду. Все клиентские "
        "счётчики совпали с числом строк в ClickHouse, ошибок и сообщений context "
        "deadline exceeded не было.",
    )
    insert_after(
        result_anchor,
        "Пиковое наблюдаемое потребление памяти ClickHouse составило около "
        "2,67 ГиБ при лимите Docker Desktop 7,75 ГиБ; Go-сервер использовал около "
        "20 МиБ. Следовательно, часовой профиль поместился в выделенную память с "
        "запасом примерно 5 ГиБ, хотя объём хранилища продолжает расти линейно.",
    )

    set_paragraph(
        doc,
        "Расхождение составило 0,101%",
        "В исходном высоком профиле расхождение составило 0,101% сигналов и "
        "0,084% событий. В часовом профиле после исправления полнота составила "
        "100% для обоих типов сообщений. Это подтверждает, что третий тест "
        "корректно выявил ошибку завершения, а четвёртый проверил её устранение.",
    )
    set_paragraph(
        doc,
        "Исправление должно включать",
        "Для устранения обнаруженной причины выполнены два точечных изменения:",
    )
    set_paragraph(
        doc,
        "1. Не применять обычный",
        "1. WebSocket-маршрут исключён из 60-секундного middleware.Timeout; "
        "ограничение сохранено для REST-запросов.",
    )
    set_paragraph(
        doc,
        "2. При нештатном ответе",
        "2. Нагрузочный клиент постоянно читает соединение и корректно "
        "обрабатывает ping/pong в длительных сессиях.",
    )
    set_paragraph(
        doc,
        "Дополнительно session_end",
        "Механизм возврата записей в буфер при произвольной ошибке репозитория "
        "пока не реализован и остаётся отдельным направлением повышения "
        "надёжности.",
    )
    set_paragraph(
        doc,
        "После изменения необходимо",
        "Регрессионный профиль продолжался один час. Критерии исправления "
        "выполнены: клиентские счётчики совпали с числом строк, а сообщения "
        "context deadline exceeded в серверном журнале отсутствовали.",
    )

    archive_table = find_table(doc, "Профиль", 4)
    append_row(
        archive_table,
        ["Часовой, 1 сессия", "81 247 397", "25 251 962", "68,9%"],
        size=10,
    )
    set_paragraph(
        doc,
        "Во всех профилях gzip",
        "Во всех профилях gzip уменьшил объём примерно на 68–69% относительно "
        "несжатого размера данных ClickHouse. Для одной сессии часового профиля "
        "81,25 МБ несжатых данных превратились в два архива общим размером "
        "25,25 МБ.",
    )
    set_paragraph(
        doc,
        "Близкий процент во всех трёх",
        "Близкий процент во всех четырёх измерениях объясняется одинаковой "
        "структурой тестовых записей. Для часового профиля в таблицу и диаграмму "
        "включена одна сессия, а не сумма десяти сессий. Для реальных данных "
        "коэффициент может отличаться, особенно при больших уникальных строках.",
    )
    archive_result = find_paragraph(doc, "Близкий процент во всех четырёх")
    archive_result = insert_after(
        archive_result,
        "Для оценки ожидания клиента архивировались две большие сессии. Сессия с "
        "1 896 850 сигналами и 75 877 событиями была обработана за 19,94 секунды. "
        "Сессия с 1 897 250 сигналами и 75 871 событием — за 29,12 секунды. В обоих "
        "случаях HTTP 200 был получен только после формирования двух файлов.",
    )
    insert_after(
        archive_result,
        "После ответа обновление статуса и DROP PARTITION выполнялись асинхронно. "
        "По журналу ClickHouse три фоновые команды заняли 8, 10 и 7 мс. Отдельного "
        "уведомления о завершении фоновой стадии нет.",
    )

    directions_table = find_table(doc, "Направление", 3)
    replace_row(
        directions_table,
        1,
        [
            "Политика timeout для WebSocket",
            "Исправлено: timeout применяется только к REST",
            "Проверено часовым профилем без потерь",
        ],
        size=11,
    )
    append_row(
        directions_table,
        [
            "Длительное архивирование",
            "Клиент ожидает 19,94–29,12 с; предел REST — 60 с",
            "Фоновое задание и endpoint состояния",
        ],
        size=10,
    )
    set_paragraph(
        doc,
        "Наиболее важным направлением",
        "Таймаут WebSocket устранён и проверен часовым профилем. Наиболее важными "
        "оставшимися направлениями являются защита буфера при ошибке записи и "
        "перевод длительной архивации в наблюдаемое фоновое задание.",
    )
    set_paragraph(
        doc,
        "1. Исключить WebSocket",
        "1. Сохранять раздельную политику timeout для WebSocket и REST и покрыть "
        "её автоматическим регрессионным тестом.",
    )
    changes_anchor = find_paragraph(doc, "8. Перейти на официальный")
    insert_after(
        changes_anchor,
        "9. Запускать длительное архивирование как фоновое задание и добавить "
        "REST-маршрут получения состояния операции.",
    )

    set_paragraph(
        doc,
        "Система успешно обработала лёгкий",
        "Система успешно обработала лёгкий и средний профили без потерь. Исходный "
        "высокий профиль достиг 4961 сигнала в секунду и выявил конфликт "
        "WebSocket с общим HTTP timeout. После точечного исправления часовой "
        "профиль сохранил 17 995 850 сигналов и 719 665 событий без расхождений "
        "при средней скорости 4999 сигналов в секунду.",
    )
    set_paragraph(
        doc,
        "Архивирование работает",
        "Архивирование уменьшает объём данных примерно на две трети. Формирование "
        "файлов выполняется синхронно и заняло 19,94–29,12 секунды, после чего "
        "статус и удаление партиций обрабатываются в фоне без отдельного "
        "уведомления. Для более крупных сессий требуется модель фоновых заданий.",
    )

    set_paragraph(
        doc,
        "ПРИЛОЖЕНИЕ 4.",
        "ПРИЛОЖЕНИЕ 4. ЖУРНАЛ ВЫСОКОГО И ЧАСОВОГО НАГРУЗОЧНЫХ ТЕСТОВ\t47",
    )
    replace_paragraph_text(
        find_heading(
            doc, "ПРИЛОЖЕНИЕ 4. ЖУРНАЛ ВЫСОКОГО НАГРУЗОЧНОГО ТЕСТА"
        ),
        "ПРИЛОЖЕНИЕ 4. ЖУРНАЛ ВЫСОКОГО И ЧАСОВОГО НАГРУЗОЧНЫХ ТЕСТОВ",
    )
    set_paragraph(
        doc,
        "Сессий: 10",
        "Высокий профиль до исправления; сессий: 10",
    )
    appendix_anchor = find_paragraph(doc, "Событий сохранено: 11890")
    for line in [
        "Часовой профиль после исправления",
        "Сессий: 10; длительность: 3600 секунд",
        "Сигналов отправлено и сохранено: 17 995 850",
        "Событий отправлено и сохранено: 719 665",
        "Средняя скорость: 4999 сигналов/с и 200 событий/с",
        "Клиентских ошибок: 0",
        "Пиковая память: ClickHouse 2,67 ГиБ; сервер около 20 МиБ",
        "Архивация одной сессии: 1 897 250 сигналов и 75 871 событие; "
        "29,12 с; gzip 25 251 962 байта",
    ]:
        appendix_anchor = insert_after(appendix_anchor, line)

    set_paragraph(
        doc,
        "Функциональные тесты подтвердили",
        "Функциональные тесты подтвердили запуск окружения, проверку API-ключа, "
        "получение данных, архивирование и полное удаление сессии. Исходный "
        "высокий профиль выявил потерю около 0,1% данных из-за отмены контекста "
        "на 60-й секунде. После исключения WebSocket из обычного HTTP timeout и "
        "обработки ping/pong часовой профиль сохранил 17 995 850 сигналов и "
        "719 665 событий без потерь.",
    )
    set_paragraph(
        doc,
        "Полученные результаты показывают",
        "Полученные результаты показывают, что система устойчиво обрабатывает "
        "около 5 тыс. сигналов в секунду в течение одного часа. ClickHouse "
        "использовал до 2,67 ГиБ памяти из 7,75 ГиБ, а сервер — около 20 МиБ.",
    )
    set_paragraph(
        doc,
        "Архивирование в JSON Lines",
        "Архивирование в JSON Lines с gzip уменьшило объём большой сессии на "
        "68,9%. Клиент синхронно ожидает формирования файлов 19,94–29,12 секунды. "
        "Обновление статуса и удаление активных партиций выполняются после ответа "
        "асинхронно, отдельного уведомления об их завершении нет.",
    )
    set_paragraph(
        doc,
        "Цель работы достигнута",
        "Цель работы достигнута на уровне программной реализации. Реализован и "
        "экспериментально проверен полный цикл данных: часовой приём, пакетная "
        "запись, чтение, архивирование большого набора и удаление. Четвёртый "
        "нагрузочный профиль подтвердил исправление обнаруженного timeout и "
        "позволил определить следующий приоритет — фоновую архивацию с "
        "наблюдаемым состоянием.",
    )

    find_heading(
        doc, "5.6. Архивирование и сжатие"
    ).paragraph_format.page_break_before = True

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    with tempfile.TemporaryDirectory(prefix="diploma_charts_") as chart_dir_name:
        chart_dir = Path(chart_dir_name)
        speed = chart_dir / "speed.png"
        completeness = chart_dir / "completeness.png"
        archive = chart_dir / "archive.png"

        categories = ["Лёгкий", "Средний", "Высокий", "1 час"]
        draw_grouped_chart(
            speed,
            (1800, 1000),
            "Фактическая скорость передачи сообщений",
            categories,
            [
                ("Сигналы, сообщений/с", [296, 1486, 4961, 4999]),
                ("События, сообщений/с", [15, 50, 198, 200]),
            ],
            5700,
            "сообщений/с",
        )
        draw_grouped_chart(
            completeness,
            (1589, 889),
            "Полнота сохранения записей в ClickHouse",
            categories,
            [
                ("Сигналы, %", [100, 100, 99.90, 100]),
                ("События, %", [100, 100, 99.92, 100]),
            ],
            105,
            "проценты",
            decimals=2,
        )
        draw_grouped_chart(
            archive,
            (1800, 1000),
            "Объём данных до и после архивирования",
            ["Лёгкий", "Средний", "Высокий", "1 час, 1 сесс."],
            [
                ("До архива, МБ", [0.27, 1.85, 12.75, 81.25]),
                ("gzip-архив, МБ", [0.09, 0.59, 3.96, 25.25]),
            ],
            92,
            "МБ",
            decimals=2,
        )

        patch_media(
            output,
            {
                "word/media/image3.png": speed,
                "word/media/image4.png": completeness,
                "word/media/image5.png": archive,
            },
        )

def normalize(text):
    return re.sub(r"[^0-9a-zа-яё]+", "", text.lower())


def finalize_document(docx_path, rendered_pdf):
    doc = Document(docx_path)
    reader = PdfReader(str(rendered_pdf))
    page_texts = [normalize(page.extract_text() or "") for page in reader.pages]
    page_count = len(reader.pages)

    headings = {}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        if not text or "Heading" not in style_name:
            continue
        normalized = normalize(text)
        page = next(
            (
                index
                for index, page_text in enumerate(page_texts[8:], start=9)
                if normalized in page_text
            ),
            None,
        )
        if page is not None:
            headings[text] = page

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "\t" not in text:
            continue
        title = text.split("\t", 1)[0]
        if title in headings:
            replace_paragraph_text(paragraph, f"{title}\t{headings[title]}")

    russian = find_paragraph(doc, "Жунусов Е.Е.")
    replace_paragraph_text(
        russian,
        re.sub(r"— \d+ с\.", f"— {page_count} с.", russian.text),
    )
    english = find_paragraph(doc, "Zhunusov E.E.")
    replace_paragraph_text(
        english,
        re.sub(r"— \d+ pages", f"— {page_count} pages", english.text),
    )

    doc.save(docx_path)


def main():
    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest="command", required=True)

    build = subparsers.add_parser("build")
    build.add_argument("--output", required=True)

    finalize = subparsers.add_parser("finalize")
    finalize.add_argument("--docx", required=True)
    finalize.add_argument("--pdf", required=True)

    args = parser.parse_args()
    if args.command == "build":
        build_document(Path(args.output))
    else:
        finalize_document(Path(args.docx), Path(args.pdf))


if __name__ == "__main__":
    main()
